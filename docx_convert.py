
import os
import re
import glob
import zipfile
import logging
import json
import shutil
from datetime import datetime
from email import policy
from email.parser import BytesParser
from email.header import decode_header
import tempfile
import pypdf
from docx import Document
from openai import AzureOpenAI

from rad_utils import numero_radicado_global

##########################
# FUNCIONES COMUNES
##########################

def decode_str(string):
    """
    Decodifica strings que pueden contener caracteres especiales o estar codificados
    en diferentes formatos de correo electrónico (Encoded-Word).
    """
    result = ""
    if string:
        decoded_parts = decode_header(string)
        for decoded_string, charset in decoded_parts:
            if isinstance(decoded_string, bytes):
                try:
                    result += decoded_string.decode(charset if charset else 'utf-8', errors='replace')
                except (UnicodeDecodeError, LookupError):
                    result += decoded_string.decode('utf-8', errors='replace')
            else:
                result += str(decoded_string)
    return result

def extract_text_from_pdf(pdf_path):
    """
    Extrae el texto de un archivo PDF.
    """
    try:
        text = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    except Exception as e:
        return f"Error al procesar PDF: {str(e)}"

def extract_text_from_docx(docx_path):
    """
    Extrae el texto de un archivo DOCX.
    """
    try:
        doc = Document(docx_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        return f"Error al procesar DOCX: {str(e)}"

def add_letter_paragraphs(doc, text):
    """
    Evita párrafos vacíos extra: sustituye múltiples saltos de línea por uno solo,
    y solo agrega un párrafo si la línea no está en blanco.
    """
    text = text.replace("\\n", "\n")
    text = re.sub(r'\n+', '\n', text)
    for line in text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())

##########################
# FUNCIONES PARA PROCESAR EL .eml
##########################

def get_latest_zip(downloads_path):
    """
    Busca el último archivo ZIP descargado en la carpeta 'downloads_path' y lo mueve
    a la carpeta 'downloads_path/ZIP' para su procesamiento.
    """
    zip_dir = os.path.join(downloads_path, "ZIP")
    if not os.path.exists(zip_dir):
        os.makedirs(zip_dir)
    
    # ZIPs que ya estén en la carpeta ZIP
    zip_files_in_zip_dir = glob.glob(os.path.join(zip_dir, "*.zip"))
    # ZIPs en la carpeta principal de descargas
    zip_files_in_downloads = glob.glob(os.path.join(downloads_path, "*.zip"))
    
    if zip_files_in_downloads:
        latest_zip = max(zip_files_in_downloads, key=os.path.getmtime)
        destino = os.path.join(zip_dir, os.path.basename(latest_zip))
        try:
            os.rename(latest_zip, destino)
            zip_files_in_zip_dir.append(destino)
        except Exception as e:
            raise Exception(f"Error al mover el archivo {latest_zip} a {zip_dir}: {e}")
    
    if not zip_files_in_zip_dir:
        raise FileNotFoundError(f"No se encontraron archivos ZIP en {zip_dir} ni en {downloads_path}")
    
    return max(zip_files_in_zip_dir, key=os.path.getmtime)

def extract_eml(zip_path):
    """
    Extrae el primer archivo .eml encontrado en el ZIP y lo coloca en un directorio temporal.
    """
    temp_dir = "temp_extract"
    os.makedirs(temp_dir, exist_ok=True)
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        all_files = zip_ref.namelist()
        eml_files = [f for f in all_files if f.endswith('.eml')]
        if not eml_files:
            raise FileNotFoundError("No se encontró ningún archivo .eml en el ZIP")
        zip_ref.extract(eml_files[0], temp_dir)
        return os.path.join(temp_dir, eml_files[0])

def parse_eml_with_attachments(eml_path):
    """
    Parsea el archivo .eml con BytesParser y extrae:
      - Encabezados (from, to, subject, date)
      - Cuerpo principal (text/plain)
      - Contenido de adjuntos PDF y DOCX
    Devuelve un diccionario con todos los datos.
    """
    with open(eml_path, 'rb') as fp:
        msg = BytesParser(policy=policy.default).parse(fp)
    
    # Encabezados
    from_ = decode_str(msg.get('from', 'No disponible'))
    to_ = decode_str(msg.get('to', 'No disponible'))
    subject_ = decode_str(msg.get('subject', 'No disponible'))
    date_ = decode_str(msg.get('date', 'No disponible'))
    
    # Obtener el cuerpo principal text/plain (si existe)
    body_text = []
    attachments_text = []
    
    with tempfile.TemporaryDirectory() as temp_dir:
        for part in msg.walk():
            # Ignoramos partes multipart
            if part.get_content_maintype() == 'multipart':
                continue
            
            # 1) Contenido 'text/plain' => cuerpo
            if part.get_content_type() == 'text/plain':
                try:
                    body = part.get_payload(decode=True).decode(
                        part.get_content_charset() or 'utf-8',
                        errors='replace'
                    )
                    body_text.append(body)
                except Exception as e:
                    body_text.append(f"Error al decodificar body: {str(e)}")
            
            # 2) Archivos adjuntos => PDF / DOCX
            content_disposition = str(part.get('Content-Disposition') or "")
            if 'attachment' in content_disposition:
                filename = part.get_filename()
                if filename:
                    filename_decoded = decode_str(filename)
                    filepath = os.path.join(temp_dir, filename_decoded)
                    # Guardar binario
                    with open(filepath, 'wb') as f:
                        f.write(part.get_payload(decode=True))
                    
                    # PDF
                    if filename_decoded.lower().endswith('.pdf'):
                        pdf_text = extract_text_from_pdf(filepath)
                        attachments_text.append(
                            f"\n=== [Adjunto PDF: {filename_decoded}] ===\n{pdf_text}"
                        )
                    # DOCX
                    elif filename_decoded.lower().endswith('.docx'):
                        docx_text = extract_text_from_docx(filepath)
                        attachments_text.append(
                            f"\n=== [Adjunto DOCX: {filename_decoded}] ===\n{docx_text}"
                        )
    
    combined_body = "\n".join(body_text)
    combined_attachments = "\n".join(attachments_text)
    
    return {
        "from": from_,
        "to": to_,
        "subject": subject_,
        "date": date_,
        "body": combined_body,
        "attachments": combined_attachments
    }

##########################
# CONFIGURACIÓN DEL CLIENTE AZURE OPENAI
##########################

endpoint = "https://respuestas-agiles4.openai.azure.com/"
deployment = "gpt-4oPrueba01"  # Ajusta el nombre de tu deployment
subscription_key = "CDZld9bCoXXuELG9gxIlGwNooRGixWx2qnH6xsxxP461EedloO2cJQQJ99BAACHYHv6XJ3w3AAABACOGt0o5"

client = AzureOpenAI(
    azure_endpoint=endpoint,
    api_key=subscription_key,
    api_version="2024-05-01-preview",
)

# Cargamos las instrucciones de sistema desde un archivo .md
instructions_file = r"C:\Users\SBURGOSP\Desktop\config\instrucciones.md"
with open(instructions_file, "r", encoding="utf-8") as file:
    instructions_content = file.read()

##########################
# FLUJO PRINCIPAL
##########################

def process_main():
    # Configuración de logging básico
    logging.basicConfig(level=logging.INFO)
    logger = logging.getLogger("IntegratedProcess")
    
    downloads_path = os.path.expanduser("~/Downloads")
    
    try:
        # 1. Obtener el último ZIP y extraer el .eml
        zip_path = get_latest_zip(downloads_path)
        logger.info(f"ZIP encontrado: {zip_path}")
        
        eml_path = extract_eml(zip_path)
        logger.info(f"EML extraído: {eml_path}")
        
        # 2. Parsear el EML con extracción de adjuntos
        email_data = parse_eml_with_attachments(eml_path)
        logger.info("Correo procesado correctamente (con adjuntos).")
        
    except Exception as e:
        logger.error(f"Error procesando el email: {str(e)}")
        return


  #  numero_rad = numero_radicado_global if numero_radicado_global is not None else "No disponible"

    # 3. Construir el prompt para Azure OpenAI
    user_prompt = f"""

    cuando leas el PQR primero haz una validacion de acuerdo a las instrucciones del "instructions_content" para que sepas como debes responder las cartas de respuesta.
    Debes responder en un formato JSON que sea valido que no tenga '\n\' o '\f\'
    La fecha debe ser en español 
    
{{
  "carta_firmante": "Texto de la carta de respuesta al firmante juntos con los datos del firmante en el PQR",
  "carta_empresa": "Texto de la carta de traslado a la empresa competente junto con los datos de contacto de la empresa a la que se le hace traslado (identificar bien el objeto de la PQR para identifacar a donde se realiza traslado, en este PQR de ejemplo hace referencia a un traslado por estado del alumbrado publico del municipio de Ocaña)"
}}

=== DATOS DEL CORREO ===
From: {email_data['from']}
To: {email_data['to']}
Subject: {email_data['subject']}
Date: {email_data['date']}

=== CONTENIDO DEL CORREO ===
{email_data['body']}

=== CONTENIDO DE ADJUNTOS (si hay) ===
{email_data['attachments']}
"""

    chat_prompt = [
        {
            "role": "system",
            "content": instructions_content
        },
        {
            "role": "user",
            "content": user_prompt
        }
    ]

    # 4. Llamada al modelo de Azure OpenAI
    try:
        completion = client.chat.completions.create(
            model=deployment,
            messages=chat_prompt,
            max_tokens=15000,
            temperature=0.1,
            top_p=0.95,
            frequency_penalty=0,
            presence_penalty=0,
            stop=None,
            stream=False
        )
    except Exception as e:
        logger.error(f"Error al llamar al modelo Azure OpenAI: {str(e)}")
        return


# Extraer el contenido generado por el modelo
    response_json = completion.to_json()
    response_text = json.loads(response_json)["choices"][0]["message"]["content"]

    # 5. Buscar bloque JSON en la respuesta
    first_bracket = response_text.find('{')
    last_bracket = response_text.rfind('}')
    if first_bracket != -1 and last_bracket != -1:
        json_str = response_text[first_bracket:last_bracket+1]
        logger.info(f"JSON extraído:\n{json_str}")
    else:
        logger.info("No se encontró un bloque JSON en la respuesta.")
        json_str = response_text

    # 6. Parsear el JSON con las dos cartas
    try:
        response_data = json.loads(json_str)
        carta_firmante = response_data.get("carta_firmante", "").strip()
        carta_empresa = response_data.get("carta_empresa", "").strip()
    except json.JSONDecodeError as e:
        logger.error(f"Error al parsear el JSON: {e}")
        # Extraer con regex
        carta_firmante, carta_empresa = "", ""
        m_firmante = re.search(r'"carta_firmante"\s*:\s*"(.+?)"\s*,', json_str, re.DOTALL)
        if m_firmante:
            carta_firmante = m_firmante.group(1).strip()
        m_empresa = re.search(r'"carta_empresa"\s*:\s*"(.+?)"\s*}', json_str, re.DOTALL)
        if m_empresa:
            carta_empresa = m_empresa.group(1).strip()


    # 7. Generar los documentos Word con las cartas
    if carta_firmante:
        doc_firmante = Document()
        add_letter_paragraphs(doc_firmante, carta_firmante)
        output_docx_firmante = "carta_firmante.docx"
        doc_firmante.save(output_docx_firmante)
        logger.info(f"Carta del firmante guardada en {output_docx_firmante}")
    else:
        logger.info("No se encontró contenido para la carta del firmante.")

    if carta_empresa:
        doc_empresa = Document()
        add_letter_paragraphs(doc_empresa, carta_empresa)
        output_docx_empresa = "carta_empresa.docx"
        doc_empresa.save(output_docx_empresa)
        logger.info(f"Carta de la empresa guardada en {output_docx_empresa}")

    # 8. Limpiar directorio temporal
    if os.path.exists("temp_extract"):
        shutil.rmtree("temp_extract")
        logger.info("Directorio temp_extract eliminado.")

if __name__ == "__main__":
    process_main()




















#El número de radicado para este caso es: {numero_radicado_global}